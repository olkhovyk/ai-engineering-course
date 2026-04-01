"""Generate realistic healthcare industry documents for the ingestion pipeline.

Creates properly formatted, openable PDF, DOCX, XLSX, and HTML files
that simulate real healthcare enterprise documents.
"""

from pathlib import Path
from datetime import datetime, timedelta
import random

SAMPLES_DIR = Path(__file__).parent.parent / "samples" / "healthcare"

random.seed(42)  # reproducible samples


# ---------------------------------------------------------------------------
# 1. Large multi-page PDF: Clinical Trial Report
# ---------------------------------------------------------------------------
def create_clinical_trial_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, ListFlowable, ListItem,
    )
    from reportlab.lib import colors

    path = SAMPLES_DIR / "clinical_trial_report.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontSize=18)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14)
    body = styles["Normal"]

    elements = []

    # Cover page
    elements.append(Spacer(1, 5*cm))
    elements.append(Paragraph("CLINICAL TRIAL REPORT", title_style))
    elements.append(Spacer(1, 1*cm))
    elements.append(Paragraph("Study ID: HCR-2025-0847", body))
    elements.append(Paragraph("Protocol: Phase III, Randomized, Double-Blind", body))
    elements.append(Paragraph("Sponsor: MedPharm International Inc.", body))
    elements.append(Paragraph("Date: December 15, 2025", body))
    elements.append(Paragraph("Status: CONFIDENTIAL", body))
    elements.append(PageBreak())

    # Table of Contents
    elements.append(Paragraph("Table of Contents", h2_style))
    toc = [
        "1. Executive Summary", "2. Study Design and Methodology",
        "3. Patient Demographics", "4. Efficacy Results",
        "5. Safety Profile", "6. Adverse Events Summary",
        "7. Laboratory Findings", "8. Statistical Analysis",
        "9. Discussion", "10. Conclusions and Recommendations",
    ]
    for item in toc:
        elements.append(Paragraph(item, body))
    elements.append(PageBreak())

    # 1. Executive Summary
    elements.append(Paragraph("1. Executive Summary", h2_style))
    elements.append(Paragraph(
        "This Phase III clinical trial evaluated the efficacy and safety of Compound XR-7821 "
        "(200mg oral, twice daily) versus placebo in 1,247 patients with moderate-to-severe "
        "Type 2 Diabetes Mellitus over a 52-week treatment period. The primary endpoint was "
        "change from baseline in HbA1c at Week 24. Secondary endpoints included fasting "
        "plasma glucose (FPG), body weight change, and proportion of patients achieving "
        "HbA1c < 7.0%.", body))
    elements.append(Spacer(1, 0.5*cm))
    elements.append(Paragraph(
        "Key findings: Compound XR-7821 demonstrated statistically significant superiority "
        "over placebo in reducing HbA1c (least-squares mean difference: -1.2%, 95% CI: "
        "-1.4 to -1.0, p < 0.001). The safety profile was consistent with the known "
        "pharmacological class, with no new safety signals identified.", body))
    elements.append(PageBreak())

    # 2. Study Design
    elements.append(Paragraph("2. Study Design and Methodology", h2_style))
    elements.append(Paragraph(
        "This was a multicenter, randomized, double-blind, placebo-controlled, parallel-group "
        "study conducted at 87 sites across 12 countries (United States, Canada, United Kingdom, "
        "Germany, France, Spain, Italy, Japan, South Korea, Australia, Brazil, and India). "
        "The study was conducted in accordance with ICH-GCP guidelines and the Declaration "
        "of Helsinki. All patients provided written informed consent prior to enrollment.", body))
    elements.append(Spacer(1, 0.5*cm))

    # Inclusion criteria
    elements.append(Paragraph("Inclusion Criteria:", ParagraphStyle("Bold", parent=body, fontName="Helvetica-Bold")))
    criteria = [
        "Age 18-75 years, diagnosed with Type 2 Diabetes Mellitus for >= 6 months",
        "HbA1c 7.5% - 10.5% at screening",
        "BMI 25-40 kg/m²",
        "Stable dose of metformin (>= 1500 mg/day) for >= 8 weeks",
        "eGFR >= 60 mL/min/1.73m²",
    ]
    for c in criteria:
        elements.append(Paragraph(f"• {c}", body))
    elements.append(PageBreak())

    # 3. Patient Demographics
    elements.append(Paragraph("3. Patient Demographics", h2_style))
    demo_data = [
        ["Characteristic", "XR-7821 (N=624)", "Placebo (N=623)", "Total (N=1247)"],
        ["Age, mean (SD)", "54.3 (10.2)", "53.8 (10.5)", "54.1 (10.3)"],
        ["Female, n (%)", "298 (47.8%)", "305 (49.0%)", "603 (48.4%)"],
        ["White", "412 (66.0%)", "408 (65.5%)", "820 (65.8%)"],
        ["Black/African American", "87 (13.9%)", "92 (14.8%)", "179 (14.3%)"],
        ["Asian", "98 (15.7%)", "95 (15.2%)", "193 (15.5%)"],
        ["Hispanic/Latino", "27 (4.3%)", "28 (4.5%)", "55 (4.4%)"],
        ["BMI, mean (SD)", "31.2 (4.1)", "31.5 (4.3)", "31.4 (4.2)"],
        ["HbA1c baseline, mean (SD)", "8.4% (0.9)", "8.3% (0.8)", "8.4% (0.9)"],
        ["Diabetes duration, years", "8.7 (5.2)", "8.5 (5.0)", "8.6 (5.1)"],
    ]
    t = Table(demo_data, colWidths=[5*cm, 3.5*cm, 3.5*cm, 3.5*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#ecf0f1")]),
    ]))
    elements.append(t)
    elements.append(PageBreak())

    # 4. Efficacy
    elements.append(Paragraph("4. Efficacy Results", h2_style))
    elements.append(Paragraph(
        "The primary endpoint analysis demonstrated that XR-7821 was superior to placebo "
        "in reducing HbA1c from baseline to Week 24. The adjusted mean change from baseline "
        "was -1.6% for XR-7821 compared to -0.4% for placebo, with a treatment difference "
        "of -1.2% (95% CI: -1.4 to -1.0; p < 0.001).", body))
    elements.append(Spacer(1, 0.5*cm))

    efficacy_data = [
        ["Endpoint", "XR-7821", "Placebo", "Difference (95% CI)", "p-value"],
        ["HbA1c change (%)", "-1.6", "-0.4", "-1.2 (-1.4, -1.0)", "< 0.001"],
        ["FPG change (mg/dL)", "-42.3", "-8.1", "-34.2 (-39.1, -29.3)", "< 0.001"],
        ["Body weight (kg)", "-2.8", "+0.3", "-3.1 (-3.7, -2.5)", "< 0.001"],
        ["HbA1c < 7.0% (%)", "52.1%", "18.3%", "OR: 4.8 (3.7, 6.2)", "< 0.001"],
    ]
    t2 = Table(efficacy_data, colWidths=[3.5*cm, 2.5*cm, 2.5*cm, 4*cm, 2.5*cm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(t2)
    elements.append(PageBreak())

    # 5. Safety + Adverse Events
    elements.append(Paragraph("5. Safety Profile", h2_style))
    elements.append(Paragraph(
        "Overall, XR-7821 was well-tolerated. The most common adverse events (AEs) were "
        "gastrointestinal in nature, consistent with the mechanism of action. Serious adverse "
        "events (SAEs) occurred in 4.2% of XR-7821 patients and 3.5% of placebo patients. "
        "No deaths were attributed to study drug. Three cases of pancreatitis were reported "
        "(2 in XR-7821, 1 in placebo), all resolved without sequelae.", body))
    elements.append(Spacer(1, 0.5*cm))

    ae_data = [
        ["Adverse Event", "XR-7821 n (%)", "Placebo n (%)"],
        ["Nausea", "124 (19.9%)", "32 (5.1%)"],
        ["Diarrhea", "87 (13.9%)", "28 (4.5%)"],
        ["Headache", "62 (9.9%)", "58 (9.3%)"],
        ["Hypoglycemia", "45 (7.2%)", "12 (1.9%)"],
        ["Upper respiratory infection", "41 (6.6%)", "38 (6.1%)"],
        ["Injection site reaction", "0 (0%)", "0 (0%)"],
        ["Pancreatitis", "2 (0.3%)", "1 (0.2%)"],
        ["Any SAE", "26 (4.2%)", "22 (3.5%)"],
    ]
    t3 = Table(ae_data, colWidths=[5*cm, 4*cm, 4*cm])
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#c0392b")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(t3)
    elements.append(PageBreak())

    # 6-10: More sections with substantial text
    for section_num, (title, content) in enumerate([
        ("6. Adverse Events — Detailed Narrative",
         "Patient 0847-0312 (58-year-old male, XR-7821 group): Experienced Grade 2 nausea "
         "beginning on Day 14 of treatment. The event was managed with antiemetic therapy and "
         "resolved within 5 days without dose modification. The investigator assessed this event "
         "as probably related to study drug. The patient continued treatment and completed the "
         "study without recurrence. " * 5),
        ("7. Laboratory Findings",
         "Comprehensive metabolic panel, complete blood count, and urinalysis were performed "
         "at screening, baseline, and Weeks 4, 12, 24, 36, and 52. Hepatic transaminases "
         "(ALT, AST) showed no clinically significant changes from baseline in either group. "
         "Lipid panels demonstrated favorable changes in the XR-7821 group: LDL cholesterol "
         "decreased by 8.3% (vs 1.2% placebo), HDL increased by 4.1% (vs 0.8% placebo), "
         "and triglycerides decreased by 12.7% (vs 2.3% placebo). Renal function markers "
         "(eGFR, serum creatinine) remained stable throughout the study. " * 3),
        ("8. Statistical Analysis",
         "The primary analysis was performed on the modified intention-to-treat (mITT) "
         "population using a mixed-model repeated measures (MMRM) approach. The model included "
         "treatment group, visit, treatment-by-visit interaction, baseline HbA1c (continuous), "
         "and region as covariates. Missing data were handled under the missing-at-random (MAR) "
         "assumption. Sensitivity analyses included pattern-mixture models and tipping-point "
         "analyses. All sensitivity analyses supported the primary analysis conclusions. " * 3),
        ("9. Discussion",
         "The results of this Phase III study demonstrate that XR-7821 provides clinically "
         "meaningful and statistically significant reductions in HbA1c compared to placebo "
         "in patients with inadequately controlled Type 2 Diabetes on metformin monotherapy. "
         "The magnitude of HbA1c reduction (-1.2% vs placebo) exceeds the FDA guidance threshold "
         "for clinical significance (0.3-0.4%) and is consistent with results from Phase II "
         "studies. The additional benefits of weight loss and improved lipid profile suggest "
         "potential cardiovascular risk reduction, although a dedicated cardiovascular outcomes "
         "trial (CVOT) would be required to confirm this hypothesis. " * 3),
        ("10. Conclusions and Recommendations",
         "Based on the totality of evidence from this Phase III clinical trial, XR-7821 at "
         "a dose of 200mg twice daily demonstrates a favorable benefit-risk profile for the "
         "treatment of Type 2 Diabetes Mellitus. We recommend proceeding with regulatory "
         "submission for marketing authorization. Key supporting data include: statistically "
         "significant HbA1c reduction, clinically meaningful weight loss, favorable lipid "
         "effects, and an acceptable safety profile consistent with the drug class. Post-marketing "
         "surveillance should include monitoring for pancreatitis and long-term cardiovascular "
         "outcomes assessment."),
    ], start=6):
        elements.append(Paragraph(title, h2_style))
        elements.append(Paragraph(content, body))
        elements.append(PageBreak())

    doc.build(elements)
    size_kb = path.stat().st_size / 1024
    print(f"  Created: {path.name} ({size_kb:.0f} KB, multi-page clinical trial report)")


# ---------------------------------------------------------------------------
# 2. Large PDF: Patient Discharge Summary (unstructured mixed content)
# ---------------------------------------------------------------------------
def create_discharge_summary_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    path = SAMPLES_DIR / "discharge_summary_patient_4821.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    bold = ParagraphStyle("Bold", parent=styles["Normal"], fontName="Helvetica-Bold")
    body = styles["Normal"]
    h2 = styles["Heading2"]

    elements = []

    # Header block (unstructured - mix of key-value pairs, free text, lists)
    elements.append(Paragraph("PATIENT DISCHARGE SUMMARY", styles["Title"]))
    elements.append(Spacer(1, 0.3*cm))

    header_data = [
        ["Patient Name:", "John Robert Anderson", "MRN:", "MRN-2025-48210"],
        ["Date of Birth:", "March 15, 1958", "Age:", "67 years"],
        ["Admission Date:", "November 28, 2025", "Discharge Date:", "December 8, 2025"],
        ["Attending:", "Dr. Sarah Chen, MD", "Service:", "Internal Medicine"],
        ["Insurance:", "Medicare Part A", "Room:", "4-West, Bed 2"],
    ]
    t = Table(header_data, colWidths=[3*cm, 5*cm, 3*cm, 5*cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.5*cm))

    # Diagnosis
    elements.append(Paragraph("DIAGNOSES", h2))
    elements.append(Paragraph("Primary: Community-Acquired Pneumonia (ICD-10: J18.9)", bold))
    elements.append(Paragraph("Secondary:", bold))
    secondary = [
        "Type 2 Diabetes Mellitus, uncontrolled (E11.65)",
        "Hypertension, essential (I10)",
        "Chronic Kidney Disease, Stage 3a (N18.31)",
        "Hyperlipidemia (E78.5)",
        "Obesity, BMI 34.2 (E66.01)",
    ]
    for dx in secondary:
        elements.append(Paragraph(f"  • {dx}", body))

    # Hospital course (free text narrative — unstructured)
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph("HOSPITAL COURSE", h2))
    elements.append(Paragraph(
        "Mr. Anderson is a 67-year-old male with a past medical history significant for "
        "Type 2 Diabetes, Hypertension, CKD Stage 3a, and Hyperlipidemia who presented to "
        "the Emergency Department on November 28, 2025, with a 4-day history of productive "
        "cough with yellow-green sputum, fever (Tmax 102.4°F), shortness of breath on "
        "exertion, and general malaise. Chest X-ray revealed right lower lobe consolidation "
        "consistent with pneumonia. Initial labs showed WBC 14,200/μL with left shift, "
        "CRP 12.8 mg/dL, procalcitonin 2.3 ng/mL, and creatinine 1.8 mg/dL (baseline 1.4).", body))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph(
        "The patient was started on IV ceftriaxone 1g daily and oral azithromycin 500mg daily "
        "per CAP guidelines. Blood cultures drawn prior to antibiotics grew no organisms after "
        "5 days. Sputum culture grew Streptococcus pneumoniae sensitive to ceftriaxone. "
        "Supplemental oxygen was provided via nasal cannula at 2-4 L/min for the first 72 hours. "
        "The patient's diabetes was managed with sliding scale insulin; his home metformin was "
        "held due to acute kidney injury. Creatinine peaked at 2.1 on hospital day 3, then "
        "trended down to 1.6 by discharge with IV fluid hydration.", body))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph(
        "By hospital day 7, the patient was afebrile for >48 hours, WBC normalized to "
        "8,400/μL, CRP decreased to 2.1, oxygen saturation was 96% on room air, and he "
        "was tolerating oral intake well. He was transitioned to oral levofloxacin to "
        "complete a 7-day total antibiotic course. Physical therapy assessed his mobility "
        "and recommended home exercises. Social work arranged a follow-up appointment.", body))

    # Medications at discharge (structured list)
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph("DISCHARGE MEDICATIONS", h2))
    meds = [
        ("Levofloxacin", "500mg PO daily", "3 days remaining"),
        ("Metformin", "1000mg PO BID", "Resume when creatinine < 1.5"),
        ("Lisinopril", "20mg PO daily", "Continue home dose"),
        ("Atorvastatin", "40mg PO at bedtime", "Continue home dose"),
        ("Amlodipine", "5mg PO daily", "Continue home dose"),
        ("Albuterol inhaler", "2 puffs Q4H PRN", "New prescription"),
        ("Acetaminophen", "650mg PO Q6H PRN", "For fever or pain"),
    ]
    med_data = [["Medication", "Dose/Route", "Instructions"]]
    for med in meds:
        med_data.append(list(med))
    t2 = Table(med_data, colWidths=[4*cm, 4*cm, 5*cm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(t2)

    # Follow-up
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph("FOLLOW-UP INSTRUCTIONS", h2))
    elements.append(Paragraph(
        "1. PCP follow-up with Dr. Michael Torres within 7 days of discharge.<br/>"
        "2. Repeat chest X-ray in 6 weeks to confirm resolution.<br/>"
        "3. BMP and creatinine check in 5 days (lab order placed).<br/>"
        "4. Return to ED if: fever > 101°F, worsening shortness of breath, chest pain, "
        "hemoptysis, or inability to take oral medications.<br/>"
        "5. Pneumococcal vaccination (PCV20) recommended at follow-up visit.", body))

    doc.build(elements)
    size_kb = path.stat().st_size / 1024
    print(f"  Created: {path.name} ({size_kb:.0f} KB, discharge summary with mixed content)")


# ---------------------------------------------------------------------------
# 3. DOCX: HIPAA Compliance Policy
# ---------------------------------------------------------------------------
def create_hipaa_policy_docx():
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    path = SAMPLES_DIR / "hipaa_compliance_policy.docx"
    doc = Document()

    # Title
    title = doc.add_heading("HIPAA Compliance Policy and Procedures", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("Effective Date: January 1, 2026")
    doc.add_paragraph("Version: 3.2 | Last Reviewed: December 2025")
    doc.add_paragraph("Approved by: Chief Compliance Officer, General Counsel")
    doc.add_paragraph("")

    sections = [
        ("1. Purpose and Scope",
         "This policy establishes the requirements for protecting Protected Health Information "
         "(PHI) in compliance with the Health Insurance Portability and Accountability Act of "
         "1996 (HIPAA), the Health Information Technology for Economic and Clinical Health "
         "(HITECH) Act, and applicable state privacy laws. This policy applies to all workforce "
         "members including employees, contractors, volunteers, and business associates who "
         "create, receive, maintain, or transmit PHI in any form (electronic, paper, or oral)."),
        ("2. Definitions",
         "Protected Health Information (PHI): Any individually identifiable health information "
         "that is created, received, maintained, or transmitted by a covered entity. This includes "
         "demographic data, medical histories, test results, insurance information, and any other "
         "information that can be used to identify a patient.\n\n"
         "Electronic Protected Health Information (ePHI): PHI in electronic form, including data "
         "at rest (stored), data in motion (transmitted), and data in use (being processed).\n\n"
         "Minimum Necessary Standard: The principle that workforce members should only access, "
         "use, or disclose the minimum amount of PHI necessary to accomplish the intended purpose."),
        ("3. Administrative Safeguards",
         "3.1 Security Officer: The organization shall designate a Security Officer responsible "
         "for developing and implementing security policies.\n\n"
         "3.2 Workforce Training: All workforce members must complete HIPAA privacy and security "
         "training within 30 days of hire and annually thereafter. Training records must be "
         "maintained for 6 years.\n\n"
         "3.3 Access Controls: Role-based access control (RBAC) must be implemented for all "
         "systems containing ePHI. Access must be reviewed quarterly and terminated within "
         "24 hours upon workforce member separation.\n\n"
         "3.4 Risk Assessment: A comprehensive security risk assessment must be conducted "
         "annually, with remediation plans for identified vulnerabilities."),
        ("4. Technical Safeguards",
         "4.1 Encryption: All ePHI must be encrypted using AES-256 at rest and TLS 1.3 in "
         "transit. Mobile devices must use full-disk encryption.\n\n"
         "4.2 Audit Controls: All access to ePHI must be logged, including user ID, timestamp, "
         "and action performed. Logs must be retained for 6 years and reviewed monthly.\n\n"
         "4.3 Authentication: Multi-factor authentication (MFA) is required for all systems "
         "containing ePHI. Password requirements: minimum 12 characters, complexity rules, "
         "90-day rotation, no reuse of last 12 passwords.\n\n"
         "4.4 Automatic Logoff: Workstations must lock after 5 minutes of inactivity. "
         "Clinical application sessions must timeout after 15 minutes."),
        ("5. Breach Notification Procedures",
         "5.1 In the event of a breach of unsecured PHI affecting 500 or more individuals, "
         "the organization must notify: (a) affected individuals within 60 days, (b) the HHS "
         "Secretary within 60 days, and (c) prominent media outlets serving the affected area.\n\n"
         "5.2 For breaches affecting fewer than 500 individuals, notification to HHS must be "
         "submitted annually (by March 1 for breaches in the prior calendar year).\n\n"
         "5.3 All suspected breaches must be reported internally within 24 hours to the "
         "Privacy Officer via the incident reporting system."),
        ("6. Business Associate Agreements",
         "All business associates who create, receive, maintain, or transmit PHI on behalf "
         "of the organization must execute a Business Associate Agreement (BAA) prior to "
         "any PHI disclosure. BAAs must include: permitted uses and disclosures, safeguard "
         "requirements, breach notification obligations, and termination provisions. "
         "The compliance office maintains a registry of all active BAAs."),
        ("7. Sanctions Policy",
         "Violations of this policy may result in disciplinary action up to and including "
         "termination of employment. Additionally, HIPAA violations may result in civil "
         "penalties ($100-$50,000 per violation, up to $1.5 million per year per violation "
         "category) and criminal penalties (up to $250,000 fine and 10 years imprisonment "
         "for intentional violations)."),
    ]

    for heading, content in sections:
        doc.add_heading(heading, level=1)
        for para in content.split("\n\n"):
            doc.add_paragraph(para)

    doc.save(str(path))
    size_kb = path.stat().st_size / 1024
    print(f"  Created: {path.name} ({size_kb:.0f} KB, HIPAA compliance policy)")


# ---------------------------------------------------------------------------
# 4. XLSX: Patient Lab Results (large, structured)
# ---------------------------------------------------------------------------
def create_lab_results_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, numbers

    path = SAMPLES_DIR / "patient_lab_results_q4.xlsx"
    wb = Workbook()

    # Sheet 1: CBC results
    ws1 = wb.active
    ws1.title = "CBC Results"
    headers = ["Patient ID", "Name", "DOB", "Collection Date", "WBC", "RBC", "Hemoglobin",
               "Hematocrit", "Platelets", "MCV", "MCH", "MCHC", "RDW", "Status"]
    ws1.append(headers)
    for cell in ws1[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="2C3E50", fill_type="solid")

    first_names = ["James", "Mary", "Robert", "Patricia", "John", "Jennifer", "Michael",
                   "Linda", "David", "Elizabeth", "William", "Barbara", "Richard", "Susan"]
    last_names = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia", "Miller",
                  "Davis", "Rodriguez", "Martinez", "Hernandez", "Lopez", "Wilson", "Anderson"]

    for i in range(1, 501):
        status = random.choice(["Normal", "Normal", "Normal", "Abnormal High", "Abnormal Low", "Critical"])
        ws1.append([
            f"PT-{i:05d}",
            f"{random.choice(first_names)} {random.choice(last_names)}",
            f"{random.randint(1940, 2005)}-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
            f"2025-{random.randint(10,12):02d}-{random.randint(1,28):02d}",
            round(random.uniform(3.5, 18.0), 1),   # WBC
            round(random.uniform(3.5, 6.5), 2),     # RBC
            round(random.uniform(8.0, 18.0), 1),    # Hemoglobin
            round(random.uniform(25, 55), 1),        # Hematocrit
            random.randint(100, 450),                 # Platelets
            round(random.uniform(70, 110), 1),       # MCV
            round(random.uniform(25, 35), 1),        # MCH
            round(random.uniform(30, 38), 1),        # MCHC
            round(random.uniform(11, 18), 1),        # RDW
            status,
        ])

    # Sheet 2: Metabolic Panel
    ws2 = wb.create_sheet("Metabolic Panel")
    met_headers = ["Patient ID", "Glucose", "BUN", "Creatinine", "eGFR", "Sodium",
                   "Potassium", "Chloride", "CO2", "Calcium", "Total Protein", "Albumin",
                   "Bilirubin", "ALT", "AST", "ALP"]
    ws2.append(met_headers)
    for cell in ws2[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="2C3E50", fill_type="solid")

    for i in range(1, 501):
        ws2.append([
            f"PT-{i:05d}",
            random.randint(65, 350),    # Glucose
            random.randint(6, 45),      # BUN
            round(random.uniform(0.5, 3.5), 1),  # Creatinine
            random.randint(15, 120),    # eGFR
            random.randint(130, 150),   # Sodium
            round(random.uniform(3.0, 6.0), 1),  # Potassium
            random.randint(95, 112),    # Chloride
            random.randint(18, 32),     # CO2
            round(random.uniform(7.5, 11.5), 1), # Calcium
            round(random.uniform(5.5, 8.5), 1),  # Total Protein
            round(random.uniform(2.5, 5.5), 1),  # Albumin
            round(random.uniform(0.1, 3.0), 1),  # Bilirubin
            random.randint(5, 120),     # ALT
            random.randint(5, 100),     # AST
            random.randint(30, 200),    # ALP
        ])

    wb.save(str(path))
    size_kb = path.stat().st_size / 1024
    print(f"  Created: {path.name} ({size_kb:.0f} KB, 500 patients x 2 panels)")


# ---------------------------------------------------------------------------
# 5. HTML: EHR Patient Portal Page
# ---------------------------------------------------------------------------
def create_ehr_portal_html():
    path = SAMPLES_DIR / "ehr_patient_portal.html"

    vitals_rows = ""
    base_date = datetime(2025, 10, 1)
    for i in range(30):
        d = base_date + timedelta(days=i * 3)
        vitals_rows += f"""        <tr>
            <td>{d.strftime('%Y-%m-%d')}</td>
            <td>{random.randint(110, 155)}/{random.randint(65, 95)}</td>
            <td>{random.randint(60, 95)}</td>
            <td>{round(random.uniform(97.0, 99.5), 1)}</td>
            <td>{random.randint(93, 100)}%</td>
            <td>{round(random.uniform(18, 24), 1)}</td>
        </tr>
"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <title>Patient Portal - John Anderson - MRN: 48210</title>
    <meta name="description" content="Electronic Health Record patient summary">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 20px; color: #333; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #2980b9; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th {{ background: #2c3e50; color: white; padding: 10px; text-align: left; }}
        td {{ border: 1px solid #ddd; padding: 8px; }}
        tr:nth-child(even) {{ background: #f2f2f2; }}
        .alert {{ background: #e74c3c; color: white; padding: 10px; border-radius: 5px; margin: 10px 0; }}
        .info {{ background: #3498db; color: white; padding: 10px; border-radius: 5px; margin: 10px 0; }}
        .section {{ margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }}
    </style>
</head>
<body>
    <h1>Electronic Health Record - Patient Summary</h1>

    <div class="section">
        <h2>Patient Information</h2>
        <table>
            <tr><th>Field</th><th>Value</th></tr>
            <tr><td>Full Name</td><td>John Robert Anderson</td></tr>
            <tr><td>MRN</td><td>48210</td></tr>
            <tr><td>Date of Birth</td><td>March 15, 1958 (Age: 67)</td></tr>
            <tr><td>Gender</td><td>Male</td></tr>
            <tr><td>Blood Type</td><td>O+</td></tr>
            <tr><td>Primary Care</td><td>Dr. Michael Torres, MD</td></tr>
            <tr><td>Insurance</td><td>Medicare Part A & B</td></tr>
            <tr><td>Allergies</td><td><strong>Penicillin</strong> (anaphylaxis), <strong>Sulfa drugs</strong> (rash)</td></tr>
        </table>
    </div>

    <div class="alert">
        <strong>ALLERGY ALERT:</strong> Patient has documented anaphylactic reaction to Penicillin.
        Avoid all beta-lactam antibiotics without allergy testing.
    </div>

    <div class="section">
        <h2>Active Problem List</h2>
        <table>
            <tr><th>Condition</th><th>ICD-10</th><th>Onset</th><th>Status</th></tr>
            <tr><td>Type 2 Diabetes Mellitus</td><td>E11.65</td><td>2012</td><td>Active, uncontrolled</td></tr>
            <tr><td>Essential Hypertension</td><td>I10</td><td>2008</td><td>Active, controlled</td></tr>
            <tr><td>CKD Stage 3a</td><td>N18.31</td><td>2020</td><td>Active, stable</td></tr>
            <tr><td>Hyperlipidemia</td><td>E78.5</td><td>2010</td><td>Active, on treatment</td></tr>
            <tr><td>Obesity (BMI 34.2)</td><td>E66.01</td><td>2015</td><td>Active</td></tr>
            <tr><td>Community-Acquired Pneumonia</td><td>J18.9</td><td>Nov 2025</td><td>Resolved</td></tr>
        </table>
    </div>

    <div class="section">
        <h2>Current Medications</h2>
        <table>
            <tr><th>Medication</th><th>Dose</th><th>Route</th><th>Frequency</th><th>Prescriber</th></tr>
            <tr><td>Metformin</td><td>1000mg</td><td>Oral</td><td>BID</td><td>Dr. Torres</td></tr>
            <tr><td>Lisinopril</td><td>20mg</td><td>Oral</td><td>Daily</td><td>Dr. Torres</td></tr>
            <tr><td>Atorvastatin</td><td>40mg</td><td>Oral</td><td>QHS</td><td>Dr. Torres</td></tr>
            <tr><td>Amlodipine</td><td>5mg</td><td>Oral</td><td>Daily</td><td>Dr. Torres</td></tr>
            <tr><td>Aspirin</td><td>81mg</td><td>Oral</td><td>Daily</td><td>Dr. Torres</td></tr>
            <tr><td>Levofloxacin</td><td>500mg</td><td>Oral</td><td>Daily x3 days</td><td>Dr. Chen</td></tr>
        </table>
    </div>

    <div class="section">
        <h2>Vital Signs History (Last 90 Days)</h2>
        <table>
            <tr><th>Date</th><th>BP (mmHg)</th><th>HR (bpm)</th><th>Temp (°F)</th><th>SpO2</th><th>BMI</th></tr>
{vitals_rows}
        </table>
    </div>

    <div class="section">
        <h2>Recent Lab Results (December 2025)</h2>
        <table>
            <tr><th>Test</th><th>Result</th><th>Reference Range</th><th>Flag</th></tr>
            <tr><td>HbA1c</td><td>8.4%</td><td>< 7.0%</td><td style="color:red"><strong>HIGH</strong></td></tr>
            <tr><td>Fasting Glucose</td><td>186 mg/dL</td><td>70-100</td><td style="color:red"><strong>HIGH</strong></td></tr>
            <tr><td>Creatinine</td><td>1.6 mg/dL</td><td>0.7-1.3</td><td style="color:red"><strong>HIGH</strong></td></tr>
            <tr><td>eGFR</td><td>48 mL/min</td><td>> 60</td><td style="color:red"><strong>LOW</strong></td></tr>
            <tr><td>Total Cholesterol</td><td>195 mg/dL</td><td>< 200</td><td>Normal</td></tr>
            <tr><td>LDL</td><td>112 mg/dL</td><td>< 100</td><td style="color:orange">Borderline</td></tr>
            <tr><td>HDL</td><td>42 mg/dL</td><td>> 40</td><td>Normal</td></tr>
            <tr><td>WBC</td><td>8,400/μL</td><td>4,500-11,000</td><td>Normal</td></tr>
            <tr><td>Hemoglobin</td><td>13.2 g/dL</td><td>13.5-17.5</td><td style="color:orange">Borderline Low</td></tr>
        </table>
    </div>

    <div class="info">
        <strong>CARE COORDINATION NOTE:</strong> Patient discharged Dec 8, 2025 following
        pneumonia hospitalization. Follow-up chest X-ray due Jan 19, 2026. Creatinine
        trending down — recheck in 5 days. Consider endocrinology referral for diabetes
        optimization.
    </div>
</body>
</html>"""

    path.write_text(html)
    size_kb = path.stat().st_size / 1024
    print(f"  Created: {path.name} ({size_kb:.0f} KB, EHR patient portal page)")


def main():
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    print("Generating healthcare industry documents...\n")

    create_clinical_trial_pdf()
    create_discharge_summary_pdf()
    create_hipaa_policy_docx()
    create_lab_results_xlsx()
    create_ehr_portal_html()

    total = len(list(SAMPLES_DIR.iterdir()))
    print(f"\nDone! Created {total} healthcare documents in {SAMPLES_DIR}")
    print(f"\nRun: python src/main.py --input {SAMPLES_DIR}")


if __name__ == "__main__":
    main()
