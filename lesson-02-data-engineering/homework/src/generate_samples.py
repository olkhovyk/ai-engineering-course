"""Generate sample PDF, DOCX, and HTML files for testing the pipeline."""

import sys
from pathlib import Path

SAMPLES_DIR = Path(__file__).parent.parent / "samples"

SAMPLE_TEXT = (
    "Artificial intelligence (AI) is transforming how organizations process and "
    "understand documents. Modern AI pipelines require clean, structured data extracted "
    "from diverse sources such as PDFs, Word documents, and web pages.\n\n"
    "Data ingestion is the first critical step. Raw documents must be parsed, cleaned, "
    "and chunked into manageable segments before they can be fed into language models "
    "or vector databases.\n\n"
    "Optical Character Recognition (OCR) plays a vital role when dealing with scanned "
    "documents. Unlike born-digital PDFs where text can be extracted directly, scanned "
    "documents are essentially images and require OCR to convert visual content back "
    "into machine-readable text.\n\n"
    "Streaming ingestion allows pipelines to process documents as they arrive, rather "
    "than waiting for a complete batch. This is essential for real-time applications "
    "where latency matters -- for example, processing customer support tickets or "
    "monitoring news feeds.\n\n"
    "Data versioning ensures reproducibility. By tracking snapshots of processed data, "
    "teams can roll back to any previous state, compare changes between versions, and "
    "maintain audit trails for compliance."
)


def create_pdf():
    """Create a sample PDF file."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        path = SAMPLES_DIR / "sample_report.pdf"
        c = canvas.Canvas(str(path), pagesize=A4)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 780, "AI Data Engineering Report")
        c.setFont("Helvetica", 11)

        y = 750
        for line in SAMPLE_TEXT.split("\n"):
            if not line.strip():
                y -= 15
                continue
            # Wrap long lines
            words = line.split()
            current_line = ""
            for word in words:
                if len(current_line + word) > 80:
                    c.drawString(50, y, current_line)
                    y -= 15
                    current_line = word + " "
                else:
                    current_line += word + " "
            if current_line:
                c.drawString(50, y, current_line)
                y -= 15

        c.save()
        print(f"  Created: {path}")
    except ImportError:
        # Fallback: create a minimal valid PDF manually
        path = SAMPLES_DIR / "sample_report.pdf"
        pdf_content = _make_minimal_pdf("AI Data Engineering Report", SAMPLE_TEXT)
        path.write_bytes(pdf_content)
        print(f"  Created (minimal): {path}")


def _make_minimal_pdf(title: str, body: str) -> bytes:
    """Create a minimal valid PDF without external dependencies."""
    text = body.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    lines = text.split("\n")
    text_ops = ""
    for line in lines:
        text_ops += f"({line}) Tj T* "

    stream = f"""BT
/F1 12 Tf
50 750 Td
14 TL
({title}) Tj T*
T*
/F1 10 Tf
{text_ops}
ET"""
    stream_bytes = stream.encode("latin-1")

    objects = []

    # Obj 1: Catalog
    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    # Obj 2: Pages
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    # Obj 3: Page
    objects.append(
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R "
        b"/MediaBox [0 0 595 842] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    )
    # Obj 4: Stream
    objects.append(
        f"4 0 obj\n<< /Length {len(stream_bytes)} >>\nstream\n".encode()
        + stream_bytes
        + b"\nendstream\nendobj\n"
    )
    # Obj 5: Font
    objects.append(
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    )

    pdf = b"%PDF-1.4\n"
    offsets = []
    for obj in objects:
        offsets.append(len(pdf))
        pdf += obj

    xref_offset = len(pdf)
    pdf += b"xref\n"
    pdf += f"0 {len(objects) + 1}\n".encode()
    pdf += b"0000000000 65535 f \n"
    for off in offsets:
        pdf += f"{off:010d} 00000 n \n".encode()

    pdf += b"trailer\n"
    pdf += f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    pdf += b"startxref\n"
    pdf += f"{xref_offset}\n".encode()
    pdf += b"%%EOF\n"

    return pdf


def create_docx():
    """Create a sample DOCX file."""
    try:
        from docx import Document
        from docx.shared import Pt

        path = SAMPLES_DIR / "sample_document.docx"
        doc = Document()

        doc.add_heading("AI Data Engineering Report", level=1)

        for para in SAMPLE_TEXT.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Add a sample table
        doc.add_heading("Pipeline Metrics", level=2)
        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"
        headers = ["Stage", "Documents", "Avg Time (ms)"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        data = [
            ("Parsing", "1,247", "45"),
            ("Chunking", "1,247", "12"),
            ("Versioning", "1,247", "8"),
        ]
        for row_idx, (stage, docs, time_ms) in enumerate(data, 1):
            table.rows[row_idx].cells[0].text = stage
            table.rows[row_idx].cells[1].text = docs
            table.rows[row_idx].cells[2].text = time_ms

        doc.save(str(path))
        print(f"  Created: {path}")
    except ImportError:
        print("  Skipped DOCX (install python-docx: pip install python-docx)")


def create_html():
    """Create a sample HTML file."""
    path = SAMPLES_DIR / "sample_page.html"

    paragraphs_html = ""
    for para in SAMPLE_TEXT.split("\n\n"):
        if para.strip():
            paragraphs_html += f"    <p>{para.strip()}</p>\n"

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="description" content="AI Data Engineering course material on document ingestion">
    <title>AI Data Engineering — Document Ingestion</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #2c3e50; }}
    </style>
    <script>
        // This script should be stripped by the parser
        console.log("This should not appear in extracted text");
    </script>
</head>
<body>
    <h1>AI Data Engineering — Document Ingestion</h1>
{paragraphs_html}
    <h2>Useful Links</h2>
    <ul>
        <li><a href="https://example.com/ocr">OCR Best Practices</a></li>
        <li><a href="https://example.com/streaming">Streaming Architecture Guide</a></li>
        <li><a href="https://example.com/versioning">Data Versioning Patterns</a></li>
    </ul>
</body>
</html>"""

    path.write_text(html)
    print(f"  Created: {path}")


def create_xlsx():
    """Create a sample Excel file."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        path = SAMPLES_DIR / "sample_metrics.xlsx"
        wb = Workbook()

        # Sheet 1: Pipeline metrics
        ws1 = wb.active
        ws1.title = "Pipeline Metrics"
        ws1.append(["Stage", "Documents", "Avg Time (ms)", "Success Rate"])
        ws1["A1"].font = Font(bold=True)
        data = [
            ("Ingestion", 1247, 45, "98.5%"),
            ("Parsing", 1247, 120, "95.2%"),
            ("Chunking", 1187, 12, "100%"),
            ("Embedding", 1187, 340, "99.1%"),
            ("Indexing", 1176, 55, "99.8%"),
        ]
        for row in data:
            ws1.append(row)

        # Sheet 2: Document inventory
        ws2 = wb.create_sheet("Document Inventory")
        ws2.append(["ID", "Filename", "Type", "Size (KB)", "Status"])
        for i in range(1, 21):
            ws2.append([
                i,
                f"document_{i:03d}.pdf",
                "PDF" if i % 3 else "DOCX",
                50 + i * 12,
                "Processed" if i % 4 else "Failed",
            ])

        wb.save(str(path))
        print(f"  Created: {path}")
    except ImportError:
        print("  Skipped XLSX (install openpyxl: pip install openpyxl)")


def main():
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    print("Generating sample documents...")
    create_pdf()
    create_docx()
    create_html()
    create_xlsx()
    print("Done! Check the samples/ directory.")


if __name__ == "__main__":
    main()
